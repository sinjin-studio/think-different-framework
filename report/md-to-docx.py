#!/usr/bin/env python3
"""
Convert Think Different Framework markdown presentations to branded .docx files.

Usage: python3 md-to-docx.py input.md output.docx
"""

import sys
import re
import os


def check_dependencies():
    try:
        import docx
        return True
    except ImportError:
        print("python-docx not installed. Install with: pip3 install python-docx", file=sys.stderr)
        print("Markdown output is still available.", file=sys.stderr)
        return False


def parse_markdown(text):
    """Parse the presentation markdown into metadata, The Line, and typed sections."""
    metadata = {}
    the_lines = []
    sections = {}  # keyed by type: 'insight', 'brief', 'manifesto'

    # Parse header metadata (lines like > **Key:** Value)
    for match in re.finditer(r'>\s*\*\*(\w[\w\s]*):\*\*\s*(.+)', text):
        key = match.group(1).strip().lower()
        metadata[key] = match.group(2).strip()

    # Parse The Line section - extract numbered or plain lines
    line_section = re.search(r'## The Line\s*\n(.*?)(?=\n---|\n## )', text, re.DOTALL)
    if line_section:
        line_text = line_section.group(1).strip()
        for line in line_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            # Strip numbering (e.g. "1. The line" -> "The line")
            line = re.sub(r'^\d+\.\s*', '', line)
            # Strip wrapping quotes
            line = line.strip('"\'')
            if line:
                the_lines.append(line)

    # Split by ## headers to find typed sections
    type_sections = re.split(r'\n## (Insight|Creative Brief|Manifesto)\s*\n', text)
    # type_sections: [preamble, "Insight", content, "Creative Brief", content, ...]
    i = 1
    while i < len(type_sections) - 1:
        section_type = type_sections[i].strip().lower()
        section_content = type_sections[i + 1].strip()
        # Normalize type key
        if section_type == 'creative brief':
            section_type = 'brief'
        sections[section_type] = section_content
        i += 2

    return metadata, the_lines, sections


def parse_subsections(text):
    """Parse bold-labeled subsections from a content block."""
    subsections = []
    pattern = r'\*\*([^*]+)\*\*\s*\n(.*?)(?=\n\*\*[^*]+\*\*|\Z)'
    for match in re.finditer(pattern, text, re.DOTALL):
        title = match.group(1).strip()
        body = match.group(2).strip()
        if body:
            subsections.append((title, body))
    # If no subsections found, treat the whole text as one block
    if not subsections:
        subsections.append(('', text))
    return subsections


def create_docx(metadata, the_lines, sections, output_path):
    """Create a branded .docx from parsed content."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    GREEN = RGBColor(0x16, 0x9B, 0x62)
    ORANGE = RGBColor(0xFF, 0x82, 0x00)
    DARK = RGBColor(0x2A, 0x2A, 0x2A)
    GREY = RGBColor(0x99, 0x99, 0x99)
    LIGHT_GREY = RGBColor(0xCC, 0xCC, 0xCC)

    HEADING_FONT = 'Literata'
    BODY_FONT = 'Literata'

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = BODY_FONT
    font.size = Pt(11)
    font.color.rgb = DARK

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    def add_green_rule():
        """Add a thin green horizontal rule."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '169B62')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Title ──
    topic = metadata.get('seed', metadata.get('topic', 'Think Different'))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(topic)
    run.font.size = Pt(28)
    run.font.color.rgb = DARK
    run.font.name = HEADING_FONT
    run.bold = True

    # Subtitle metadata
    date_str = metadata.get('date', '')
    words_str = metadata.get('words', '')
    subtitle_parts = []
    if date_str:
        subtitle_parts.append(date_str)
    if words_str:
        subtitle_parts.append(f"{words_str} words")
    source_str = metadata.get('source', '')
    if source_str:
        subtitle_parts.append(source_str)

    if subtitle_parts:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sub.paragraph_format.space_after = Pt(6)
        run = sub.add_run(' / '.join(subtitle_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = GREY
        run.font.name = BODY_FONT

    add_green_rule()

    # ── The Line(s) ──
    if the_lines:
        # Section heading
        line_heading = doc.add_paragraph()
        line_heading.paragraph_format.space_before = Pt(24)
        line_heading.paragraph_format.space_after = Pt(12)
        run = line_heading.add_run('The Line')
        run.font.size = Pt(20)
        run.font.color.rgb = DARK
        run.font.name = HEADING_FONT
        run.bold = True

        for i, line_text in enumerate(the_lines):
            line_para = doc.add_paragraph()
            line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line_para.paragraph_format.space_before = Pt(8)
            line_para.paragraph_format.space_after = Pt(8)
            prefix = f'{i + 1}. ' if len(the_lines) > 1 else ''
            run = line_para.add_run(f'{prefix}"{line_text}"')
            run.font.size = Pt(16)
            run.font.color.rgb = ORANGE
            run.font.name = HEADING_FONT
            run.italic = True
            run.bold = True

        doc.add_paragraph()  # spacer
        add_green_rule()

    # ── Render each section type ──
    def render_subsections(content, is_first_type=False):
        """Render bold-labeled subsections from a content block."""
        subsections = parse_subsections(content)
        for title, body in subsections:
            if title:
                heading = doc.add_paragraph()
                heading.paragraph_format.space_before = Pt(18)
                heading.paragraph_format.space_after = Pt(6)
                run = heading.add_run(title)
                run.font.size = Pt(14)
                run.font.color.rgb = GREEN
                run.font.name = HEADING_FONT
                run.bold = True

            # Split body into paragraphs
            # Handle the inline Line blockquote (skip it, already rendered above)
            paragraphs = body.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue
                # Skip inline line blockquotes
                if para_text.startswith('> **"') or para_text.startswith('---'):
                    continue
                # Clean up any remaining markdown
                para_text = re.sub(r'>\s*\*\*".*?"\*\*', '', para_text).strip()
                if not para_text:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = Pt(16)
                run = p.add_run(para_text)
                run.font.size = Pt(11)
                run.font.color.rgb = DARK
                run.font.name = BODY_FONT

    type_labels = {
        'insight': 'Insight',
        'brief': 'Creative Brief',
        'manifesto': 'Manifesto'
    }

    is_first = True
    for type_key in ['insight', 'brief', 'manifesto']:
        if type_key not in sections:
            continue

        if not is_first:
            add_green_rule()

        # Type header
        type_heading = doc.add_paragraph()
        type_heading.paragraph_format.space_before = Pt(24)
        type_heading.paragraph_format.space_after = Pt(12)
        run = type_heading.add_run(type_labels[type_key])
        run.font.size = Pt(20)
        run.font.color.rgb = DARK
        run.font.name = HEADING_FONT
        run.bold = True

        render_subsections(sections[type_key], is_first)
        is_first = False

    # ── Footer ──
    add_green_rule()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(12)
    run = footer.add_run('Document prepared using the Think Different Framework by Sinjin Studio')
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GREY
    run.font.name = BODY_FONT

    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_before = Pt(2)
    run = links.add_run('sinjin.studio')
    run.font.size = Pt(8)
    run.font.color.rgb = GREEN
    run.font.name = BODY_FONT
    run = links.add_run('  |  ')
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GREY
    run = links.add_run('npmjs.com/package/@sinjin/think-different-framework')
    run.font.size = Pt(8)
    run.font.color.rgb = GREEN
    run.font.name = BODY_FONT

    doc.save(output_path)


def main():
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} input.md output.docx", file=sys.stderr)
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    if not check_dependencies():
        sys.exit(1)

    if not os.path.exists(input_path):
        print(f"Error: input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    with open(input_path, 'r') as f:
        text = f.read()

    metadata, the_lines, sections = parse_markdown(text)

    if not sections and not the_lines:
        print("Warning: no sections found in markdown", file=sys.stderr)
        sys.exit(1)

    create_docx(metadata, the_lines, sections, output_path)
    print(f"  Created: {output_path}")


if __name__ == '__main__':
    main()
